# run_analysis.py — V2: YAML config, diagnostics, imputation, HTML+DOCX report
from __future__ import annotations
import argparse
from pathlib import Path
import datetime as _dt

import pandas as pd
import yaml
from jinja2 import Environment, FileSystemLoader
from docx import Document
from docx.shared import Inches

from snippets.data_io import read_clean, set_categorical_ref
from snippets.logistic_regression import fit_logistic, or_table
from snippets.survival_analysis import km_fit_plot, cox_fit, hr_table
from snippets.diagnostics import (
    logistic_diagnostics, save_roc_plot, save_calibration_plot, cox_ph_test_table
)
from snippets.imputation import impute_covariates

def parse_args():
    ap = argparse.ArgumentParser()
    ap.add_argument('--config', required=False, help='Path to YAML config (preferred)')
    # CLI fallback (ignored if --config is provided)
    ap.add_argument('--data')
    ap.add_argument('--outcome')
    ap.add_argument('--time')
    ap.add_argument('--status')
    ap.add_argument('--covars', nargs='+')
    ap.add_argument('--group')
    ap.add_argument('--ref_group')
    ap.add_argument('--cluster')
    return ap.parse_args()

def load_config(args):
    cfg = {
        'data': 'data/analysis_dataset.csv',
        'outcome': 'outcome',
        'time': 'time',
        'status': 'status',
        'covars': ['age','sex','group','bmi'],
        'group': 'group',
        'ref_group': 'control',
        'cluster': None,
        'imputation': { 'method': 'none', 'iterative_max_iter': 10 },
        'outputs': {
            'or_table_csv': 'outputs/logistic_or_table.csv',
            'hr_table_csv': 'outputs/cox_hr_table.csv',
            'ph_table_csv': 'outputs/cox_ph_test.csv',
            'km_plot': 'outputs/km_plot.png',
            'roc_plot': 'outputs/roc_curve.png',
            'calibration_plot': 'outputs/calibration_plot.png',
            'report_html': 'outputs/report.html',
            'report_docx': 'outputs/report.docx',
        },
        'report': {
            'title': 'Project Analysis Report',
            'author': 'Analyst',
            'institution': None,
            'include_tables': True,
            'include_km_plot': True,
            'include_docx': True,
        }
    }
    if args.config:
        with open(args.config, 'r', encoding='utf-8') as f:
            loaded = yaml.safe_load(f) or {}
        def merge(d, u):
            for k, v in u.items():
                if isinstance(v, dict) and isinstance(d.get(k), dict):
                    merge(d[k], v)
                else:
                    d[k] = v
            return d
        cfg = merge(cfg, loaded)
        # SOURCE OF TRUTH: ignore any CLI overrides when --config is provided
        return cfg
    # Fallback: build cfg from CLI if no --config
    if args.data: cfg['data'] = args.data
    if args.outcome: cfg['outcome'] = args.outcome
    if args.time: cfg['time'] = args.time
    if args.status: cfg['status'] = args.status
    if args.covars: cfg['covars'] = args.covars
    if args.group: cfg['group'] = args.group
    if args.ref_group: cfg['ref_group'] = args.ref_group
    if args.cluster: cfg['cluster'] = args.cluster
    return cfg

def render_html_report(cfg, or_df, hr_df, ph_df, auc, brier, c_index, km_plot_path):
    env = Environment(loader=FileSystemLoader('python/templates'))
    tpl = env.get_template('report_template.html')
    or_html = or_df.to_html(index=False, float_format=lambda x: format(x, '.3g')) if or_df is not None else ''
    hr_html = hr_df.to_html(index=False, float_format=lambda x: format(x, '.3g')) if hr_df is not None else ''
    ph_html = ph_df.to_html(index=False, float_format=lambda x: format(x, '.3g')) if ph_df is not None else ''
    html = tpl.render(
        title=cfg['report']['title'],
        author=cfg['report']['author'],
        institution=cfg['report']['institution'],
        generated=_dt.datetime.now().strftime('%Y-%m-%d %H:%M'),
        params={
            'data': cfg['data'],
            'outcome': cfg['outcome'],
            'time': cfg['time'],
            'status': cfg['status'],
            'covars': cfg['covars'],
            'group': cfg['group'],
            'ref_group': cfg['ref_group'],
        },
        imputation={'method': cfg['imputation'].get('method','none')},
        include_tables=bool(cfg['report'].get('include_tables', True)),
        include_km_plot=bool(cfg['report'].get('include_km_plot', True)),
        or_table_html=or_html,
        hr_table_html=hr_html,
        ph_table_html=ph_html,
        auc=f"{auc:.3f}",
        brier=f"{brier:.3f}",
        c_index=f"{c_index:.3f}",
        km_plot_path=km_plot_path,
        roc_plot_path=cfg['outputs']['roc_plot'],
        calibration_plot_path=cfg['outputs']['calibration_plot'],
    )
    out_path = Path(cfg['outputs']['report_html'])
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(html, encoding='utf-8')
    return str(out_path)

def render_docx_report(cfg, or_df, hr_df, ph_df, auc, brier, c_index):
    if not bool(cfg['report'].get('include_docx', True)):
        return None
    doc = Document()
    doc.add_heading(cfg['report']['title'], level=1)
    meta = doc.add_paragraph()
    meta.add_run(f"Author: {cfg['report']['author']}")
    if cfg['report'].get('institution'):
        meta.add_run(f" — {cfg['report']['institution']}")
    meta.add_run(f"\nGenerated: {_dt.datetime.now().strftime('%Y-%m-%d %H:%M')}")

    # Parameters
    doc.add_heading('Analysis Parameters', level=2)
    p = doc.add_paragraph()
    p.add_run(f"Data: {cfg['data']}\n")
    p.add_run(f"Outcome: {cfg['outcome']}\n")
    p.add_run(f"Time / Status: {cfg['time']} / {cfg['status']}\n")
    p.add_run(f"Covariates: {', '.join(cfg['covars'])}\n")
    p.add_run(f"Group (KM): {cfg['group']} (ref: {cfg['ref_group']})\n")
    p.add_run(f"Imputation: {cfg['imputation'].get('method','none')}\n")

    # Tables
    if bool(cfg['report'].get('include_tables', True)):
        doc.add_heading('Logistic Regression (Odds Ratios)', level=2)
        t1 = doc.add_table(rows=1, cols=len(or_df.columns))
        t1.style = 'Light List'
        hdr = t1.rows[0].cells
        for i, c in enumerate(or_df.columns): hdr[i].text = str(c)
        for _, row in or_df.iterrows():
            cells = t1.add_row().cells
            for i, c in enumerate(or_df.columns):
                val = row[c]
                cells[i].text = str(round(val, 3)) if isinstance(val, (int, float)) else str(val)

        doc.add_heading('Cox Proportional Hazards (Hazard Ratios)', level=2)
        t2 = doc.add_table(rows=1, cols=len(hr_df.columns))
        t2.style = 'Light List'
        hdr = t2.rows[0].cells
        for i, c in enumerate(hr_df.columns): hdr[i].text = str(c)
        for _, row in hr_df.iterrows():
            cells = t2.add_row().cells
            for i, c in enumerate(hr_df.columns):
                val = row[c]
                cells[i].text = str(round(val, 3)) if isinstance(val, (int, float)) else str(val)

    # Diagnostics
    doc.add_heading('Diagnostics — Logistic', level=2)
    doc.add_paragraph(f"AUC: {auc:.3f}  |  Brier score: {brier:.3f}")
    roc_path = cfg['outputs']['roc_plot']
    cal_path = cfg['outputs']['calibration_plot']
    if Path(roc_path).exists(): doc.add_picture(roc_path, width=Inches(5.5))
    if Path(cal_path).exists(): doc.add_picture(cal_path, width=Inches(5.5))

    doc.add_heading('Diagnostics — Cox', level=2)
    doc.add_paragraph(f"Concordance index (c-index): {c_index:.3f}")
    if ph_df is not None and not ph_df.empty:
        t3 = doc.add_table(rows=1, cols=len(ph_df.columns))
        t3.style = 'Light List'
        hdr = t3.rows[0].cells
        for i, c in enumerate(ph_df.columns): hdr[i].text = str(c)
        for _, row in ph_df.iterrows():
            cells = t3.add_row().cells
            for i, c in enumerate(ph_df.columns):
                val = row[c]
                cells[i].text = str(round(val, 3)) if isinstance(val, (int, float)) else str(val)

    out_docx = Path(cfg['outputs']['report_docx'])
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)
    return str(out_docx)

def main():
    args = parse_args()
    if not args.config:
        print('[Warning] --config not provided. Falling back to CLI parameters as configuration.')
    cfg = load_config(args)

    # Load data & basic preparation
    df = read_clean(cfg['data'])
    # Drop rows with missing target/time/status
    df = df.dropna(subset=[cfg['outcome'], cfg['time'], cfg['status']])

    # Group reference for KM
    if cfg['group'] in df.columns:
        df = set_categorical_ref(df, cfg['group'], cfg['ref_group'])

    # Impute covariates as configured
    df = impute_covariates(df, cfg['covars'], method=cfg['imputation']['method'],
                           iterative_max_iter=int(cfg['imputation'].get('iterative_max_iter', 10)))

    # Logistic model
    log_res = fit_logistic(df, cfg['outcome'], cfg['covars'], cluster=cfg.get('cluster'))
    or_df = or_table(log_res)
    or_csv = Path(cfg['outputs']['or_table_csv']); or_csv.parent.mkdir(parents=True, exist_ok=True)
    or_df.to_csv(or_csv, index=False)

    # Logistic diagnostics
    y_true = df[cfg['outcome']].astype(int)
    y_prob = log_res.predict()
    mets = logistic_diagnostics(y_true, y_prob)
    auc, brier = float(mets['auc']), float(mets['brier'])
    save_roc_plot(y_true, y_prob, cfg['outputs']['roc_plot'])
    save_calibration_plot(y_true, y_prob, cfg['outputs']['calibration_plot'])

    # Cox model
    cph = cox_fit(df, cfg['time'], cfg['status'], cfg['covars'])
    hr_df = hr_table(cph)
    hr_csv = Path(cfg['outputs']['hr_table_csv']); hr_csv.parent.mkdir(parents=True, exist_ok=True)
    hr_df.to_csv(hr_csv, index=False)

    # PH test table
    ph_df = cox_ph_test_table(cph, df[[cfg['time'], cfg['status']] + cfg['covars']].dropna(), cfg['time'], cfg['status'])
    ph_csv = Path(cfg['outputs']['ph_table_csv']); ph_csv.parent.mkdir(parents=True, exist_ok=True)
    ph_df.to_csv(ph_csv, index=False)

    # KM plot
    km_plot_path = None
    if bool(cfg['report'].get('include_km_plot', True)):
        ax = km_fit_plot(df, cfg['time'], cfg['status'], group=cfg['group'])
        km_png = Path(cfg['outputs']['km_plot']); km_png.parent.mkdir(parents=True, exist_ok=True)
        ax.figure.savefig(km_png, dpi=300, bbox_inches='tight')
        km_plot_path = str(km_png)

    # Reports
    html_path = render_html_report(cfg, or_df, hr_df, ph_df, auc, brier, float(getattr(cph, 'concordance_index_', float('nan'))), km_plot_path)
    docx_path = None
    if bool(cfg['report'].get('include_docx', True)):
        docx_path = render_docx_report(cfg, or_df, hr_df, ph_df, auc, brier, float(getattr(cph, 'concordance_index_', float('nan'))))

    print(f'Saved OR table → {or_csv}')
    print(f'Saved HR table → {hr_csv}')
    print(f'Saved PH test table → {ph_csv}')
    if km_plot_path: print(f'Saved KM plot → {km_plot_path}')
    print(f'Saved HTML report → {html_path}')
    if docx_path: print(f'Saved DOCX report → {docx_path}')

if __name__ == '__main__':
    main()
